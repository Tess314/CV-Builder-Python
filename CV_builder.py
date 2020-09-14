import tkinter as tk
from docx import Document

def word_build():
    document = Document()
    
    #DETAILS
    name = input("\nEnter your full name:\n")
    document.add_heading(name, 0)

    #EDUCATION
    ed_level = input("What is your highest level of education?\n")
    ed_school = input("Which school did you get this qualification/degree?\n")
    ed_date = input("When did you obtain this qualification/degree?\n")

    education = [ed_school + " - ", ed_level + " (Obtained " + ed_date + ")"]
    title1 = document.add_paragraph("EDUCATION")
    section1 = document.add_paragraph(education)

    #WORK EXPERIENCE
    company = input("Where was the last place you worked?\n")
    title = input("What was your job title?\n")
    start = input("When did you start working there?\n")
    end = input("When did you finish working there?\n")

    experience = [title + ", " + company + " (" + start + " - " + end + ")"]
    title2 = document.add_paragraph("\nWORK EXPERIENCE")
    section2 = document.add_paragraph(experience)

    #SKILLS AND INTERESTS
    personality = input("What's one word that describes you?\n")
    interest = input("What industry is the job you're applying for?\n")
    skill = input("What is your greatest skill?\n")
    print("\nCheck your files to see your brand new CV!")

    skills = ["I am " + personality + " with interest and experience in the " + interest + " sector.\nMy greatest skill is " + skill + " which makes me a valuable addition to any " + interest + " team."]
    title3 = document.add_paragraph("\nSKILLS AND INTERESTS")
    section3 = document.add_paragraph(skills)

    document.save("CV.docx")


#GUI
root = tk.Tk()
root.title("CV Builder")
canvas = tk.Canvas(root, width = 400, height = 300)
canvas.pack()

label = tk.Label(root, text='CV Builder')
label.config(font=('helvetica', 14))
canvas.create_window(200, 40, window=label)

button = tk.Button(text='Build in Word', command=word_build)
canvas.create_window(200, 140, window=button)

root.mainloop()
